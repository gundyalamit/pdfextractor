import streamlit as st
import tempfile
import os
import json
import pdfplumber
import fitz  # PyMuPDF
import google.generativeai as genai
from PIL import Image
import pandas as pd
from io import BytesIO
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side, Alignment

# For local development without Streamlit Cloud secrets, we can use python-dotenv
# If dotenv is available and .env exists, it will load it.
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

st.set_page_config(page_title="PDF Extractor", layout="wide")

# Configure Gemini API Key
api_key = os.getenv("GEMINI_API_KEY")
try:
    if not api_key and "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
except Exception:
    pass

if api_key:
    genai.configure(api_key=api_key)

def is_pdf_readable(pdf_path: str) -> bool:
    """Checks if the PDF has a significant amount of extractable text."""
    try:
        doc = fitz.open(pdf_path)
        text_length = 0
        for i in range(min(3, len(doc))):
            page = doc[i]
            text_length += len(page.get_text().strip())
            
        return text_length > 50
    except Exception as e:
        st.error(f"Error checking PDF: {e}")
        return False

def extract_with_pdfplumber(pdf_path: str):
    elements = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            tables = page.find_tables()
            tables = sorted(tables, key=lambda t: t.bbox[1])
            
            y_cursor = 0
            for table in tables:
                table_top = table.bbox[1]
                table_bottom = table.bbox[3]
                
                if table_top > y_cursor:
                    try:
                        cropped = page.crop((0, y_cursor, page.width, table_top))
                        text_above = cropped.extract_text()
                        if text_above and text_above.strip():
                            elements.append({"type": "text", "content": text_above.strip()})
                    except Exception:
                        pass
                
                extracted = table.extract()
                if extracted and len(extracted) >= 2:
                    rows_data = extracted
                    row_len = len(rows_data)
                    col_len = len(rows_data[0]) if row_len > 0 else 0
                    
                    table_obj = {"rows": []}
                    processed = [[False]*col_len for _ in range(row_len)]
                    
                    for r in range(row_len):
                        row_obj = {"cells": []}
                        for c in range(col_len):
                            if processed[r][c]:
                                row_obj["cells"].append({"value": "", "is_merged_continuation": True})
                                continue
                                
                            val = rows_data[r][c]
                            row_span = 1
                            col_span = 1
                            
                            if val is not None:
                                for cc in range(c + 1, col_len):
                                    if rows_data[r][cc] is None:
                                        col_span += 1
                                    else:
                                        break
                                for rr in range(r + 1, row_len):
                                    all_none = True
                                    for cc in range(c, c + col_span):
                                        if rows_data[rr][cc] is not None:
                                            all_none = False
                                            break
                                    if all_none:
                                        row_span += 1
                                    else:
                                        break
                                        
                                for rr in range(r, r + row_span):
                                    for cc in range(c, c + col_span):
                                        processed[rr][cc] = True
                            else:
                                processed[r][c] = True
                                val = ""
                                
                            cell_obj = {
                                "value": str(val).strip() if val else "",
                                "bold": r == 0,
                                "italic": False,
                                "borders": ["top", "bottom", "left", "right"],
                                "is_merged_continuation": False,
                                "col_span": col_span,
                                "row_span": row_span
                            }
                            row_obj["cells"].append(cell_obj)
                        table_obj["rows"].append(row_obj)
                    
                    if table_obj["rows"]:
                        elements.append({"type": "table", "rows": table_obj["rows"]})
                
                y_cursor = table_bottom
                
            if y_cursor < page.height:
                try:
                    cropped = page.crop((0, y_cursor, page.width, page.height))
                    text_below = cropped.extract_text()
                    if text_below and text_below.strip():
                        elements.append({"type": "text", "content": text_below.strip()})
                except Exception:
                    pass
                    
    return {
        "elements": elements,
        "method": "pdfplumber (Digital PDF)"
    }

def extract_with_gemini(pdf_path: str):
    if not api_key or api_key == "your_gemini_api_key_here":
        raise ValueError("GEMINI_API_KEY is not set. Please add it to Streamlit Secrets or .env")

    doc = fitz.open(pdf_path)
    images = []
    
    for i in range(len(doc)):
        page = doc[i]
        pix = page.get_pixmap(dpi=150)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
        
    prompt = """
    Analyze the provided pages of a document. 
    Extract all text and tables in the exact order they appear visually from top to bottom.
    
    Return ONLY a raw JSON object with this exact schema:
    {
      "elements": [
        {
          "type": "text",
          "content": "the text content"
        },
        {
          "type": "table",
          "rows": [
            {
              "cells": [
                {
                  "value": "cell text",
                  "bold": true,
                  "italic": false,
                  "col_span": 1,
                  "row_span": 1,
                  "borders": ["top", "bottom", "left", "right"]
                }
              ]
            }
          ]
        }
      ]
    }
    """
    
    model = genai.GenerativeModel('gemini-1.5-flash')
    contents = [prompt] + images
    
    try:
        response = model.generate_content(
            contents,
            generation_config={"response_mime_type": "application/json"}
        )
        result = json.loads(response.text)
        return {
            "elements": result.get("elements", []),
            "method": "Google Gemini API (Scanned PDF)"
        }
    except Exception as e:
        raise ValueError(f"Failed to extract with Gemini: {e}")

import re
ILLEGAL_CHARACTERS_RE = re.compile(r'[\000-\010]|[\013-\014]|[\016-\037]')

def generate_excel(data):
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "Extracted Document"
    
    merged_set = set()
    current_row = 1
    
    elements = data.get('elements', [])
    for el in elements:
        if el.get('type') == 'text':
            content = el.get('content', '')
            if isinstance(content, str):
                content = ILLEGAL_CHARACTERS_RE.sub('', content)
            
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    val = line.strip()
                    if val.startswith(('=', '+', '-', '@')):
                        val = "'" + val
                    ws.cell(row=current_row, column=1, value=val)
                    current_row += 1
            current_row += 1
            
        elif el.get('type') == 'table':
            rows = el.get("rows", [])
            for r_idx, row in enumerate(rows):
                cells = row.get("cells", [])
                for c_idx, cell in enumerate(cells):
                    if cell.get("is_merged_continuation"):
                        continue
                        
                    value = cell.get("value", "")
                    if isinstance(value, str):
                        value = ILLEGAL_CHARACTERS_RE.sub('', value)
                        value = value.replace('\n', ' ')
                        if value.startswith(('=', '+', '-', '@')):
                            value = "'" + value
                        
                    excel_cell = ws.cell(row=current_row + r_idx, column=c_idx + 1, value=value)
                    
                    is_bold = cell.get("bold", False)
                    is_italic = cell.get("italic", False)
                    if is_bold or is_italic:
                        excel_cell.font = Font(bold=is_bold, italic=is_italic)
                        
                    borders = cell.get("borders", [])
                    if isinstance(borders, list) and len(borders) > 0:
                        thin = Side(border_style="thin", color="000000")
                        excel_cell.border = Border(
                            top=thin if "top" in borders else None,
                            bottom=thin if "bottom" in borders else None,
                            left=thin if "left" in borders else None,
                            right=thin if "right" in borders else None
                        )
                        
                    col_span = int(cell.get("col_span", 1) or 1)
                    row_span = int(cell.get("row_span", 1) or 1)
                    if col_span > 1 or row_span > 1:
                        overlap = False
                        for rr in range(current_row + r_idx, current_row + r_idx + row_span):
                            for cc in range(c_idx + 1, c_idx + 1 + col_span):
                                if (rr, cc) in merged_set:
                                    overlap = True
                        if not overlap:
                            for rr in range(current_row + r_idx, current_row + r_idx + row_span):
                                for cc in range(c_idx + 1, c_idx + 1 + col_span):
                                    merged_set.add((rr, cc))
                            try:
                                ws.merge_cells(
                                    start_row=current_row + r_idx, start_column=c_idx + 1,
                                    end_row=current_row + r_idx + row_span - 1, end_column=c_idx + col_span
                                )
                            except Exception:
                                pass
                    
                    excel_cell.alignment = Alignment(vertical="center", horizontal="left", wrap_text=True)
            current_row += len(rows) + 1
                    
    wb.save(output)
    return output.getvalue()

def generate_word(data):
    doc = Document()
    doc.add_heading('Extracted Document Data', 0)
    
    elements = data.get('elements', [])
    for el in elements:
        if el.get('type') == 'text':
            content = el.get('content', '')
            if isinstance(content, str):
                content = ILLEGAL_CHARACTERS_RE.sub('', content)
            if content.strip():
                doc.add_paragraph(content)
        elif el.get('type') == 'table':
            rows = el.get("rows", [])
            if len(rows) > 0:
                num_cols = max([len(r.get("cells", [])) for r in rows]) if rows else 0
                if num_cols > 0:
                    docx_table = doc.add_table(rows=len(rows), cols=num_cols)
                    docx_table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows):
                        row_cells = docx_table.rows[r_idx].cells
                        for c_idx, cell in enumerate(row.get("cells", [])):
                            if c_idx < len(row_cells):
                                if cell.get("is_merged_continuation"):
                                    continue
                                    
                                value = cell.get("value", "")
                                if isinstance(value, str):
                                    value = ILLEGAL_CHARACTERS_RE.sub('', value)
                                    value = value.replace('\n', ' ')
                                    
                                docx_cell = row_cells[c_idx]
                                
                                col_span = int(cell.get("col_span", 1) or 1)
                                row_span = int(cell.get("row_span", 1) or 1)
                                
                                if col_span > 1 or row_span > 1:
                                    end_r = min(r_idx + row_span - 1, len(rows) - 1)
                                    end_c = min(c_idx + col_span - 1, num_cols - 1)
                                    if end_r > r_idx or end_c > c_idx:
                                        try:
                                            bottom_right_cell = docx_table.cell(end_r, end_c)
                                            docx_cell.merge(bottom_right_cell)
                                        except Exception:
                                            pass
                                
                                docx_cell.text = ""
                                if docx_cell.paragraphs:
                                    p = docx_cell.paragraphs[0]
                                else:
                                    p = docx_cell.add_paragraph()
                                    
                                run = p.add_run(str(value))
                                if cell.get("bold"):
                                    run.bold = True
                                if cell.get("italic"):
                                    run.italic = True
                                    
            doc.add_paragraph()
            
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

st.title("PDF Extraction Application")

# Sidebar for uploading
with st.sidebar:
    st.header("Upload PDF")
    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    
    if uploaded_file is not None:
        if st.button("Extract Data"):
            with st.spinner("Extracting..."):
                try:
                    # Save uploaded file to a temporary location
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                        temp_pdf.write(uploaded_file.getvalue())
                        temp_pdf_path = temp_pdf.name

                    # Determine extraction method
                    if is_pdf_readable(temp_pdf_path):
                        st.info("Detected as digital PDF. Using pdfplumber.")
                        data = extract_with_pdfplumber(temp_pdf_path)
                    else:
                        st.info("Detected as scanned PDF. Using Gemini API.")
                        data = extract_with_gemini(temp_pdf_path)
                        
                    data["filename"] = uploaded_file.name
                    st.success("Extraction successful!")
                    st.session_state["extraction_results"] = data
                    
                except ValueError as ve:
                    st.error(str(ve))
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                finally:
                    # Clean up temporary file
                    if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
                        os.remove(temp_pdf_path)

# Main area for previewing extracted data
if "extraction_results" in st.session_state:
    results = st.session_state["extraction_results"]
    
    st.header("Extracted Data")
    st.caption(f"Extraction Method Used: {results.get('method', 'Unknown')}")
    
    st.subheader("Document Preview")
    elements = results.get('elements', [])
    if not elements:
        st.info("No content extracted.")
    else:
        table_count = 1
        for i, el in enumerate(elements):
            if el.get('type') == 'text':
                st.text(el.get('content', ''))
            elif el.get('type') == 'table':
                st.write(f"**Table {table_count}**")
                table_count += 1
                df_data = []
                for row in el.get("rows", []):
                    df_row = []
                    for cell in row.get("cells", []):
                        df_row.append(cell.get("value", ""))
                    df_data.append(df_row)
                
                if df_data:
                    max_cols = max(len(row) for row in df_data)
                    for row in df_data:
                        row.extend([""] * (max_cols - len(row)))
                    
                    df = pd.DataFrame(df_data)
                    if len(df) > 1:
                        cols = df.iloc[0].astype(str).tolist()
                        seen = {}
                        for j, c in enumerate(cols):
                            if c in seen:
                                seen[c] += 1
                                cols[j] = f"{c}_{seen[c]}"
                            else:
                                seen[c] = 0
                        df.columns = cols
                        df = df[1:].reset_index(drop=True)
                    st.dataframe(df)
            
    st.header("Export Data")
    col1, col2 = st.columns(2)
    with col1:
        excel_data = generate_excel(results)
        st.download_button(
            label="Export to Excel",
            data=excel_data,
            file_name=f"{results.get('filename', 'extracted')}_data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    with col2:
        word_data = generate_word(results)
        st.download_button(
            label="Export to Word",
            data=word_data,
            file_name=f"{results.get('filename', 'extracted')}_data.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("Upload a PDF file from the sidebar and click 'Extract Data' to see results here.")
